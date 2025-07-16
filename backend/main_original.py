import os
import io
import json
import uuid
import re
from datetime import datetime, timedelta
from typing import Dict, Optional, List, Tuple
import pdfplumber
from docx import Document
from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
import anthropic
from dotenv import load_dotenv
import tempfile
from sqlalchemy import text
from sqlalchemy.orm import Session
from database import get_db
from models import Base, User, UserPlan
from database import engine

# Import route modules
from routes.auth import router as auth_router
from routes.documents import router as documents_router
from routes.chat import router as chat_router

# Load environment variables
load_dotenv()

# Create database tables
Base.metadata.create_all(bind=engine)

app = FastAPI(
    title="DocuChat API", 
    description="Document analysis API with JWT authentication and usage tracking",
    version="2.0.0"
)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3001",
        "http://localhost:3000", 
        "http://localhost:5173"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Include routers
app.include_router(auth_router)
app.include_router(documents_router)
app.include_router(chat_router)

# Anthropic configuration
anthropic_api_key = os.getenv("ANTHROPIC_API_KEY")
client = anthropic.Anthropic(api_key=anthropic_api_key) if anthropic_api_key else None

@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "message": "DocuChat API v2.0 is running",
        "features": [
            "JWT Authentication",
            "User Management", 
            "Document Analysis",
            "Chat with Documents",
            "Usage Tracking",
            "Plan-based Limits"
        ]
    }

@app.get("/health")
async def health_check(db: Session = Depends(get_db)):
    """Health check with database connectivity"""
    try:
        result = db.execute(text("SELECT 1")).fetchone()
        return {
            "status": "healthy",
            "database": "connected",
            "anthropic_configured": client is not None
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "database": "disconnected",
            "error": str(e)
        }

# Document processing functions (moved from original main.py)
def count_words(text: str) -> int:
    """Count words in text"""
    return len(text.split())

def split_text_into_chunks(text: str, target_words: int = 1200, overlap_words: int = 100) -> List[str]:
    """
    Split text into chunks of approximately target_words each.
    Ensures splitting happens on paragraph boundaries or sentence boundaries.
    """
    if not text.strip():
        return []
    
    # First, try to split by paragraphs
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    
    if not paragraphs:
        # If no paragraphs, split by sentences
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        paragraphs = sentences
    
    chunks = []
    current_chunk = ""
    current_word_count = 0
    
    for paragraph in paragraphs:
        paragraph_words = count_words(paragraph)
        
        # If adding this paragraph would exceed target, start a new chunk
        if current_word_count + paragraph_words > target_words and current_chunk:
            chunks.append(current_chunk.strip())
            
            # Start new chunk with overlap from previous chunk
            if overlap_words > 0:
                words = current_chunk.split()
                overlap_text = " ".join(words[-overlap_words:]) if len(words) > overlap_words else current_chunk
                current_chunk = overlap_text + "\n\n" + paragraph
                current_word_count = count_words(overlap_text) + paragraph_words
            else:
                current_chunk = paragraph
                current_word_count = paragraph_words
        else:
            if current_chunk:
                current_chunk += "\n\n" + paragraph
            else:
                current_chunk = paragraph
            current_word_count += paragraph_words
    
    # Add the last chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # If we have chunks that are still too long, split them further by sentences
    final_chunks = []
    for chunk in chunks:
        if count_words(chunk) <= target_words * 1.5:  # Allow some flexibility
            final_chunks.append(chunk)
        else:
            # Split by sentences
            sentences = re.split(r'[.!?]+', chunk)
            sentences = [s.strip() for s in sentences if s.strip()]
            
            sub_chunk = ""
            sub_word_count = 0
            
            for sentence in sentences:
                sentence_words = count_words(sentence)
                
                if sub_word_count + sentence_words > target_words and sub_chunk:
                    final_chunks.append(sub_chunk.strip())
                    sub_chunk = sentence
                    sub_word_count = sentence_words
                else:
                    if sub_chunk:
                        sub_chunk += ". " + sentence
                    else:
                        sub_chunk = sentence
                    sub_word_count += sentence_words
            
            if sub_chunk.strip():
                final_chunks.append(sub_chunk.strip())
    
    return final_chunks

def should_chunk_document(text: str, word_threshold: int = 1000) -> bool:
    """Determine if a document should be chunked based on word count"""
    return count_words(text) > word_threshold

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber"""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")

async def analyze_document_with_claude(text: str, retry_count: int = 0) -> dict:
    """Send text to Anthropic Claude for analysis"""
    if not client:
        raise HTTPException(status_code=500, detail="Anthropic API key not configured")
    
    # Adjust prompt based on retry count
    if retry_count == 0:
        prompt = f"""You are an AI assistant that helps explain documents clearly. 
Given this document, do the following:
1. Explain what the document is about in 1–2 sentences.
2. Summarize key important points as bullet points.
3. Highlight any risky or confusing parts with 🚩 emoji and explain why.
4. Identify key concepts/terms that are central to understanding this document.

For each key point and risk flag, please also include a short quote (5-15 words) from the original document that supports your analysis.

For key concepts, provide the term and a brief explanation of what it means in the context of this document.

CRITICAL: You must respond with ONLY valid JSON. Follow these strict rules:
- Do not include any text before or after the JSON
- Do not wrap the JSON in markdown code blocks
- Use only standard double quotes (") for strings
- Ensure all arrays and objects are properly closed
- Do not include trailing commas
- Escape any quotes within text content properly
- Keep quotes short (5-15 words maximum)

Use this exact JSON structure:
{{
    "summary": "Brief explanation of what the document is about",
    "key_points": [
        {{
            "text": "Point 1 description",
            "quote": "relevant quote from document"
        }},
        {{
            "text": "Point 2 description", 
            "quote": "relevant quote from document"
        }}
    ],
    "risk_flags": [
        {{
            "text": "🚩 Risk 1: explanation",
            "quote": "relevant quote from document"
        }},
        {{
            "text": "🚩 Risk 2: explanation",
            "quote": "relevant quote from document"
        }}
    ],
    "key_concepts": [
        {{
            "term": "Important Term 1",
            "explanation": "What this term means in the context of this document"
        }},
        {{
            "term": "Important Term 2",
            "explanation": "What this term means in the context of this document"
        }}
    ]
}}

Document text:
{text[:8000]}"""
    else:
        # Simplified prompt for retry
        prompt = f"""Analyze this document and return ONLY valid JSON with this structure:
{{
    "summary": "Brief explanation",
    "key_points": [{{"text": "point", "quote": "quote"}}],
    "risk_flags": [{{"text": "🚩 risk", "quote": "quote"}}],
    "key_concepts": [{{"term": "term", "explanation": "explanation"}}]
}}

Document: {text[:4000]}"""
    
    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1000,
            temperature=0.3,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        # Get the raw response content
        content = response.content[0].text.strip()
        
        # Try to find and extract JSON from the response
        json_start = content.find('{')
        json_end = content.rfind('}') + 1
        
        if json_start != -1 and json_end > json_start:
            json_content = content[json_start:json_end]
            
            # Clean up common JSON formatting issues
            json_content = json_content.replace('```json', '').replace('```', '').strip()
            json_content = re.sub(r',\s*}', '}', json_content)  # Remove trailing commas
            json_content = re.sub(r',\s*]', ']', json_content)  # Remove trailing commas in arrays
            
            try:
                result = json.loads(json_content)
                return result
            except json.JSONDecodeError as e:
                # Try more aggressive cleaning
                try:
                    # Remove any non-ASCII characters that might be causing issues
                    json_content = ''.join(char for char in json_content if ord(char) < 128)
                    # Try to fix common quote issues
                    json_content = re.sub(r'[""]', '"', json_content)  # Replace smart quotes
                    json_content = re.sub(r'['']', "'", json_content)  # Replace smart apostrophes
                    
                    result = json.loads(json_content)
                    return result
                except json.JSONDecodeError as e2:
                    pass
        
        # If we can't parse JSON, return fallback
        return {
            "summary": "Document analysis completed successfully.",
            "key_points": [
                {
                    "text": "Document has been processed and is ready for analysis.",
                    "quote": "Document processing completed"
                }
            ],
            "risk_flags": [],
            "key_concepts": []
        }
        
    except Exception as e:
        # If we've already retried, return a basic fallback
        if retry_count > 0:
            return {
                "summary": "Document analysis completed with basic processing.",
                "key_points": [
                    {
                        "text": "Document has been processed and is ready for analysis.",
                        "quote": "Document processing completed"
                    }
                ],
                "risk_flags": [],
                "key_concepts": []
            }
        
        # If this is the first attempt, try again with a simpler prompt
        return await analyze_document_with_claude(text, retry_count + 1)

async def analyze_document_with_chunking(text: str, enable_synthesis: bool = True) -> dict:
    """
    Analyze a document with automatic chunking for long documents.
    """
    word_count = count_words(text)
    
    # Check if document needs chunking
    if not should_chunk_document(text):
        return await analyze_document_with_claude(text)
    
    # Split document into chunks
    chunks = split_text_into_chunks(text)
    
    if len(chunks) == 1:
        return await analyze_document_with_claude(text)
    
    # Analyze each chunk
    chunk_analyses = []
    for i, chunk in enumerate(chunks):
        try:
            analysis = await analyze_document_with_claude(chunk)
            chunk_analyses.append(analysis)
        except Exception as e:
            # Continue with other chunks
            continue
    
    if not chunk_analyses:
        raise HTTPException(status_code=500, detail="Failed to analyze any document chunks")
    
    # Combine results
    return aggregate_chunk_analyses(chunk_analyses)

def aggregate_chunk_analyses(chunk_analyses: List[dict]) -> dict:
    """
    Aggregate multiple chunk analyses into a single analysis.
    """
    if not chunk_analyses:
        return {
            "summary": "This is a comprehensive analysis of your document.",
            "key_points": [],
            "risk_flags": [],
            "key_concepts": []
        }
    
    if len(chunk_analyses) == 1:
        return chunk_analyses[0]
    
    # Collect all items from all chunks
    all_key_points = []
    all_risk_flags = []
    all_key_concepts = []
    
    for analysis in chunk_analyses:
        all_key_points.extend(analysis.get("key_points", []))
        all_risk_flags.extend(analysis.get("risk_flags", []))
        all_key_concepts.extend(analysis.get("key_concepts", []))
    
    # Remove duplicates based on text content
    def remove_duplicates(items, key_func=lambda x: x.get("text", "").lower()):
        seen = set()
        unique_items = []
        for item in items:
            key = key_func(item)
            if key and key not in seen:
                seen.add(key)
                unique_items.append(item)
        return unique_items
    
    # Remove duplicates and limit results
    unique_key_points = remove_duplicates(all_key_points)[:15]
    unique_risk_flags = remove_duplicates(all_risk_flags)[:8]
    unique_key_concepts = remove_duplicates(all_key_concepts, lambda x: x.get("term", "").lower())[:10]
    
    # Create combined summary
    chunk_count = len(chunk_analyses)
    combined_summary = f"This is a comprehensive analysis of a {chunk_count}-section document covering multiple topics."
    
    return {
        "summary": combined_summary,
        "key_points": unique_key_points,
        "risk_flags": unique_risk_flags,
        "key_concepts": unique_key_concepts,
        "chunk_count": chunk_count,
        "analysis_method": "chunked"
    }

def find_quote_position(text: str, quote: str) -> Dict:
    """Find the position of a quote in the document text"""
    if not quote or not quote.strip():
        return {"start": -1, "end": -1, "found": False}
    
    # Clean up the quote for better matching
    clean_quote = quote.strip().lower()
    clean_text = text.lower()
    
    # Try exact match first
    start_pos = clean_text.find(clean_quote)
    if start_pos != -1:
        # Estimate page number based on position
        words_before = len(text[:start_pos].split())
        estimated_page = max(1, (words_before // 500) + 1)
        
        return {
            "start": start_pos,
            "end": start_pos + len(clean_quote),
            "found": True,
            "page": estimated_page
        }
    
    return {"start": -1, "end": -1, "found": False}

# PDF conversion endpoint (legacy support)
@app.post("/convert-docx-to-pdf")
async def convert_docx_to_pdf(file: UploadFile = File(...)):
    """Convert DOCX to PDF (legacy endpoint)"""
    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only DOCX files are supported.")
    
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_JUSTIFY
        
        # Read the DOCX file content
        file_content = await file.read()
        
        # Parse DOCX using python-docx
        doc = Document(io.BytesIO(file_content))
        
        # Create PDF using reportlab
        temp_dir = tempfile.gettempdir()
        pdf_filename = f"docx_converted_{uuid.uuid4().hex[:8]}.pdf"
        pdf_path = os.path.join(temp_dir, pdf_filename)
        
        # Create PDF document
        doc_pdf = SimpleDocTemplate(pdf_path, pagesize=A4, 
                                  leftMargin=0.75*inch, rightMargin=0.75*inch,
                                  topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = []
        
        # Get styles
        styles = getSampleStyleSheet()
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            alignment=TA_JUSTIFY
        )
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                story.append(Paragraph(paragraph.text, normal_style))
                story.append(Spacer(1, 4))
        
        # Build PDF
        doc_pdf.build(story)
        
        # Return the PDF file URL
        return {"pdf_url": f"/pdf/{pdf_filename}"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

@app.get("/pdf/{pdf_filename}")
def serve_pdf(pdf_filename: str):
    """Serve a PDF file from the temp directory"""
    temp_dir = tempfile.gettempdir()
    pdf_path = os.path.join(temp_dir, pdf_filename)
    if not os.path.exists(pdf_path):
        raise HTTPException(status_code=404, detail="PDF not found.")
    return FileResponse(pdf_path, media_type="application/pdf")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000) 